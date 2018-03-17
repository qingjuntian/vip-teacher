#!/usr/bin/python
#coding:utf-8

__author__ = 'qingjun'

import sys
from pypinyin import pinyin, lazy_pinyin, Style

from docx import Document

from docx.shared import Inches

class Pinyin_transfer(object):
    def __init__(self, file_in, file_out):
        self.title = u"小学五年级第一单元看拼音写汉字"
        self.space_num_per_char=8
        self.file_in = file_in
        self.file_out = file_out
        self.read_words()
        pass

    def read_words(self):
        document = Document(self.file_in)
        self.words = []

        for p in document.paragraphs:
            self.words.extend(p.text.split(' '))

class Table_transfer(Pinyin_transfer):

    def __init__(self, file_in, file_out):
        super(Table_transfer, self).__init__(file_in, file_out)

    def transfer(self):
        doc = Document()
        doc.add_heading(self.title, 0)

        row_words = self.fetch_next_row()

        while len(row_words) > 0:
            table = doc.add_table(rows=2, cols=len(row_words))
            for i in range(len(row_words)):
                py = pinyin(row_words[i])
                # table.rows[0].cells[i].
                table.rows[0].cells[i].text = u' '.join(w[0] for w in py)
                table.rows[1].cells[i].text = u'()'
            row_words = self.fetch_next_row()

        doc.save(self.file_out)
        pass

    def fetch_next_row(self):
        row_words = []
        row_word_num = 0
        index = 0
        while index < len(self.words) :
            if row_word_num > 10:
                break
            row_words.append(self.words[index])
            row_word_num += len(self.words[index])
            index += 1

        if index <= len(self.words) :
            self.words = self.words[index:]

        return row_words

        pass

    pass

class Paragraph_transfer(Pinyin_transfer):
    def __init__(self, file_in, file_out):
        super(Paragraph_transfer, self).__init__(file_in, file_out)
        #Pinyin_transfer.__init__(self, file_in, file_out)
        pass

    def transfer(self):

        doc = Document()
        doc.add_heading(self.title, 0)
        line_words = 0
        line_pinyin=''
        line_to_input=''
        char_num = 0
        first = True
        for word in self.words:
            char_num += len(word)
            py = pinyin(word)
            word_pinyin = u' '.join(w[0] for w in py)
            word_pinyin_number = len(word_pinyin)

            append_whitespace_number = (len(word) * self.space_num_per_char - word_pinyin_number * 3 / 2 + 8 / len(word)) / 2

            line_pinyin += u' ' * append_whitespace_number + word_pinyin + u' ' * append_whitespace_number

            line_to_input += u'' if first else u'  '
            line_to_input += u'(' + (u' '* len(py) * self.space_num_per_char) + ')'

            line_words += 1
            if first: first = False

            if char_num >= 14:
                #print repr(line)
                doc.add_paragraph(line_pinyin)
                doc
                doc.add_paragraph(line_to_input)
                line_pinyin = ''
                line_to_input = ''
                char_num = 0
                first = True
        doc.add_paragraph(line_pinyin)
        doc.add_paragraph(line_to_input)
        doc.save(self.file_out)


if __name__ == "__main__":
    transfer = Table_transfer(sys.argv[1], sys.argv[2])
    transfer.transfer()


#
# document = Document()
#
# document.add_heading('Document Title', 0)
#
# p = document.add_paragraph('A plain paragraph having some ')
# p.add_run('bold').bold = True
# p.add_run(' and some ')
# p.add_run('italic.').italic = True
#
# document.add_heading('Heading, level 1', level=1)
# document.add_paragraph('Intense quote', style='IntenseQuote')
#
# document.add_paragraph(
#     'first item in unordered list', style='ListBullet'
# )
# document.add_paragraph(
#     'first item in ordered list', style='ListNumber'
# )
#
# #document.add_picture('monty-truth.png', width=Inches(1.25))
#
# document.add_page_break()

# document.save(u'/Users/qingjun/workspace/vipteacher/new.docx')


